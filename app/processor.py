import os
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import pytesseract

class FileProcessor:
    """Handles text extraction and preprocessing for various file formats."""

    @staticmethod
    def extract_text(file_path):
        """Extract text from a file based on its extension.

        Returns:
            dict: A dictionary containing 'text' (extracted text) and 'metadata' (file details).
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_name = os.path.basename(file_path)
        ext = os.path.splitext(file_name)[1].lower()
        file_size = os.path.getsize(file_path)
        
        metadata = {
            "name": file_name,
            "path": file_path,
            "extension": ext,
            "size_bytes": file_size,
            "type": "Unknown"
        }

        # Check file extension and parse accordingly
        if ext == ".pdf":
            text = FileProcessor._parse_pdf(file_path)
            metadata["type"] = "PDF Document"
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            text = FileProcessor._parse_image(file_path)
            metadata["type"] = "Image File"
        elif ext == ".docx":
            text = FileProcessor._parse_docx(file_path)
            metadata["type"] = "Word Document"
        elif ext in [".txt", ".md", ".json", ".xml", ".yaml", ".yml"]:
            text = FileProcessor._parse_text(file_path)
            metadata["type"] = "Text Document"
        elif ext in [".csv", ".xlsx", ".xls"]:
            text = FileProcessor._parse_spreadsheet(file_path)
            metadata["type"] = "Spreadsheet"
        elif ext in [".py", ".js", ".html", ".css", ".java", ".c", ".cpp", ".h", ".cs", ".go", ".rs", ".sh", ".bat", ".sql"]:
            text = FileProcessor._parse_text(file_path)
            metadata["type"] = "Source Code File"
        else:
            # Fallback to plain text reading
            try:
                text = FileProcessor._parse_text(file_path)
                metadata["type"] = "Text/Data Document (Fallback)"
            except Exception:
                raise ValueError(f"Unsupported file format: {ext}")

        return {"text": text, "metadata": metadata}

    @staticmethod
    def _parse_pdf(file_path):
        """Extract text from a PDF file using PyMuPDF."""
        text_content = []
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text_content.append(page.get_text())
            doc.close()
            return "\n".join(text_content)
        except Exception as e:
            raise RuntimeError(f"Error reading PDF file: {str(e)}")

    @staticmethod
    def _parse_docx(file_path):
        """Extract text from a Word document (.docx)."""
        try:
            doc = Document(file_path)
            text_content = [paragraph.text for paragraph in doc.paragraphs]
            # Handle tables too
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_text))
            return "\n".join(text_content)
        except Exception as e:
            raise RuntimeError(f"Error reading Word document: {str(e)}")

    @staticmethod
    def _parse_text(file_path):
        """Read text from a plain text, markdown, or code file."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            raise RuntimeError(f"Error reading text file: {str(e)}")

    @staticmethod
    def _parse_spreadsheet(file_path):
        """Read spreadsheets (CSV/Excel) and format as clean text/markdown representation."""
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".csv":
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Convert first 100 rows to markdown/string to avoid overflowing LLM context
            rows_limit = 100
            truncated = len(df) > rows_limit
            sub_df = df.head(rows_limit)
            
            md_str = sub_df.to_markdown(index=False)
            info_str = f"Spreadsheet summary: {df.shape[0]} rows, {df.shape[1]} columns.\n"
            if truncated:
                info_str += f"(Showing first {rows_limit} rows)\n"
            
            return info_str + md_str
        except Exception as e:
            raise RuntimeError(f"Error reading spreadsheet: {str(e)}")

    @staticmethod
    def _parse_image(file_path):
        """Extract text from an image using PyTesseract OCR."""
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            if not text.strip():
                return "[No text detected in image via OCR]"
            return text
        except Exception as e:
            # Return custom message so vision models can handle it instead of breaking completely
            return f"[OCR Error: Tesseract might not be installed or configured correctly: {str(e)}]"
